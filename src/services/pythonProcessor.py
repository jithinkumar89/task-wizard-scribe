
import os
import re
import sys
import docx
import pandas as pd
from docx.document import Document as DocxDocument
from docx.oxml.shape import CT_Picture
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image
from io import BytesIO
import base64
import zipfile
import json
import logging
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Color
import random

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def extract_tasks_from_word(docx_path, assembly_id, assembly_name, figure_start_range, figure_end_range):
    """
    Extract tasks from a Word document that contains a table.
    
    Args:
        docx_path: Path to the Word document
        assembly_id: Assembly sequence ID to prefix task numbers
        assembly_name: Name of the assembly for task descriptions
        figure_start_range: Starting number for figure references
        figure_end_range: Ending number for figure references
    
    Returns:
        Tuple of (tasks_df, images_dict)
    """
    try:
        # Load the document
        doc = docx.Document(docx_path)
        
        # Initialize variables
        tasks = []
        images = {}
        image_map = {}  # Map to track figure references to image_ids
        image_task_mapping = {}  # Map to track which tasks reference which images
        
        # Find tables in the document
        if len(doc.tables) == 0:
            logger.warning("No tables found in the document, using paragraph-based extraction")
            return extract_tasks_from_paragraphs(doc, assembly_id, assembly_name, figure_start_range, figure_end_range)
        
        # First, try to find the most likely task table
        potential_tables = []
        for idx, table in enumerate(doc.tables):
            if len(table.rows) <= 1:  # Skip tables that are too small
                continue
            
            # Score this table based on how likely it is to be a task table
            score = 0
            
            # Check if first row is likely a header
            header_row = table.rows[0]
            header_terms = ['sl no', 'sl.no', 'sl. no', 'serial no', 'task no', 'job details', 
                           'activity', 'description', 'step', 'procedure']
            
            header_cell_texts = [cell.text.lower().strip() for cell in header_row.cells]
            
            for term in header_terms:
                if any(term in text for text in header_cell_texts):
                    score += 10
            
            # Check if subsequent rows look like tasks
            task_number_patterns = [
                r'^\s*(\d+)\s*$',              # Just a number
                r'^\s*(\d+)\.\s*$',            # Number followed by a period
                r'^\s*(\d+)\)\s*$',            # Number followed by a parenthesis
                r'^\s*[Tt]ask\s*(\d+)\s*$',    # "Task" followed by a number
                r'^\s*[Ss]tep\s*(\d+)\s*$',    # "Step" followed by a number
                r'^\s*[Ss]l\.?\s*[Nn]o\.?\s*(\d+)\s*$'  # "Sl No" followed by a number
            ]
            
            # Check first few content rows (skip header)
            task_numbers_found = 0
            for row_idx in range(1, min(6, len(table.rows))):
                if row_idx < len(table.rows):
                    row = table.rows[row_idx]
                    if row.cells and len(row.cells) > 0:
                        first_cell = row.cells[0].text.strip()
                        for pattern in task_number_patterns:
                            if re.match(pattern, first_cell):
                                task_numbers_found += 1
                                break
            
            score += task_numbers_found * 5
            
            # Score based on table size
            if len(table.rows) > 5:
                score += min(len(table.rows), 20)  # Bonus for larger tables, up to 20 rows
            
            potential_tables.append((idx, table, score))
        
        # Sort tables by score, highest first
        potential_tables.sort(key=lambda x: x[2], reverse=True)
        
        # Process each table in order of score until we find tasks
        for table_idx, table, score in potential_tables:
            logger.info(f"Processing table {table_idx} with score {score}")
            
            # Check if first row is likely a header
            header_row = table.rows[0]
            is_header = False
            
            # Check if first row contains header-like text
            header_terms = ['sl no', 'sl.no', 'sl. no', 'serial no', 'task no', 'job details', 
                           'activity', 'description', 'step', 'procedure']
            header_cell_texts = [cell.text.lower().strip() for cell in header_row.cells]
            
            for term in header_terms:
                if any(term in text for text in header_cell_texts):
                    is_header = True
                    logger.info(f"Detected header row: {' | '.join([cell.text for cell in header_row.cells])}")
                    break
            
            # Determine column indices for task number and task description
            task_num_col = 0  # Default: first column is task number
            desc_col = 1      # Default: second column is description
            
            if is_header:
                # Try to find which column contains task numbers and which contains descriptions
                for idx, cell_text in enumerate(header_cell_texts):
                    cell_lower = cell_text.lower()
                    if any(term in cell_lower for term in ['sl no', 'task no', 'step', '#', 'sl.', 'no.', 'item']):
                        task_num_col = idx
                    elif any(term in cell_lower for term in ['details', 'description', 'activity', 'action']):
                        desc_col = idx
            
            # Process each row of the table (skipping header if exists)
            for i, row in enumerate(table.rows):
                # Skip if it's the header row we identified
                if i == 0 and is_header:
                    continue
                    
                # Skip if it's likely another header row
                if any(term in cell.text.lower().strip() for cell in row.cells 
                      for term in header_terms if len(cell.text.strip()) > 0):
                    continue
                
                # Extract data from cells
                try:
                    if len(row.cells) <= max(task_num_col, desc_col):
                        continue  # Skip rows with insufficient cells
                        
                    sl_no_cell = row.cells[task_num_col].text.strip()
                    
                    # Try to extract task number using various patterns
                    task_number_patterns = [
                        r'^\s*(\d+)\s*$',              # Just a number
                        r'^\s*(\d+)\.\s*$',            # Number followed by a period
                        r'^\s*(\d+)\)\s*$',            # Number followed by a parenthesis
                        r'^\s*[Tt]ask\s*(\d+)\s*$',    # "Task" followed by a number
                        r'^\s*[Ss]tep\s*(\d+)\s*$',    # "Step" followed by a number
                        r'^\s*[Ss]l\.?\s*[Nn]o\.?\s*(\d+)\s*$'  # "Sl No" followed by a number
                    ]
                    
                    sl_no = None
                    for pattern in task_number_patterns:
                        match = re.search(pattern, sl_no_cell)
                        if match:
                            sl_no = int(match.group(1))
                            break
                    
                    if sl_no is None:
                        # If no match found using patterns, try to extract any number
                        match = re.search(r'\d+', sl_no_cell)
                        if match:
                            sl_no = int(match.group(0))
                        else:
                            # Last resort: use the row number as the task number
                            sl_no = i
                    
                    # Generate task number in the format assembly_id.0.XXX
                    task_number = f"{assembly_id}.0.{str(sl_no).zfill(3)}"
                    
                    # Extract job details
                    job_details = ""
                    if len(row.cells) > desc_col:
                        job_details_cell = row.cells[desc_col]
                        
                        # Process all paragraphs in the job details cell
                        for paragraph in job_details_cell.paragraphs:
                            job_details += paragraph.text + "\n"
                            
                            # Extract all figure references from the text
                            extract_figure_references(paragraph.text, task_number, image_task_mapping, image_map, 
                                                     figure_start_range, figure_end_range, assembly_id)
                    
                    # Create task entry with attachment information
                    task = {
                        'task_no': task_number,
                        'type': 'Operation',
                        'eta_sec': '',
                        'description': assembly_name,  # Use the provided assembly name as description
                        'activity': job_details.strip(),
                        'specification': '',
                        'attachment': ''  # Will be populated after processing all tasks and images
                    }
                    
                    tasks.append(task)
                    
                except Exception as e:
                    logger.error(f"Error processing row {i} in table {table_idx}: {e}")
                    continue
            
            # If we found tasks in this table, no need to check others
            if len(tasks) > 0:
                logger.info(f"Found {len(tasks)} tasks in table {table_idx}")
                break
        
        # If no tasks were found from tables, try paragraph-based extraction
        if len(tasks) == 0:
            logger.info("No tasks extracted from tables, trying paragraph-based extraction")
            return extract_tasks_from_paragraphs(doc, assembly_id, assembly_name, figure_start_range, figure_end_range)
        
        # Extract images from document
        all_images = extract_all_images_from_document(doc)
        logger.info(f"Found {len(all_images)} images in document")
        
        # Create final image dictionary with proper IDs
        for fig_num, image_id in image_map.items():
            # Find an image to associate with this figure number
            image_idx = min(fig_num - figure_start_range, len(all_images) - 1) if len(all_images) > 0 else -1
            if 0 <= image_idx < len(all_images):
                images[image_id] = {
                    'data': all_images[image_idx]['data'],
                    'extension': all_images[image_idx]['extension'],
                    'figure_number': fig_num
                }
        
        # If we have more images than figure references, add them with sequential IDs
        for i in range(len(image_map), len(all_images)):
            next_num = figure_start_range + i if figure_start_range > 0 else (i + 1)
            image_id = f"{assembly_id}-0-{str(next_num).zfill(3)}"
            images[image_id] = {
                'data': all_images[i]['data'],
                'extension': all_images[i]['extension'],
                'figure_number': next_num
            }
        
        # Now update each task with its attachment information (comma-separated image IDs)
        for task in tasks:
            task_num = task['task_no']
            if task_num in image_task_mapping:
                task['attachment'] = ', '.join(sorted(list(image_task_mapping[task_num])))
        
        # If we still have no task-image mappings but have both tasks and images,
        # distribute images evenly among tasks
        if len(image_task_mapping) == 0 and len(tasks) > 0 and len(images) > 0:
            logger.info("No explicit figure references found, distributing images among tasks")
            
            # For each task, assign nearby images based on relative position
            task_count = len(tasks)
            image_count = len(images)
            image_keys = list(images.keys())
            
            # Distribute images proportionally
            for i, task in enumerate(tasks):
                # Calculate which images should go to this task
                start_idx = int((i * image_count) / task_count)
                end_idx = int(((i + 1) * image_count) / task_count)
                
                if start_idx < end_idx:  # Ensure task gets at least one image
                    task_images = image_keys[start_idx:end_idx]
                    task['attachment'] = ', '.join(task_images)
        
        # Create DataFrame
        if tasks:
            df = pd.DataFrame(tasks)
            return df, images
        else:
            logger.warning("No tasks extracted from the document")
            return None, {}
            
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        return None, {}

def extract_tasks_from_paragraphs(doc, assembly_id, assembly_name, figure_start_range, figure_end_range):
    """
    Extract tasks from paragraphs when table extraction fails
    """
    try:
        tasks = []
        images = {}
        image_map = {}
        image_task_mapping = {}
        
        paragraphs = doc.paragraphs
        current_task = None
        current_task_text = ""
        task_index = 0
        
        # Define patterns to identify task headers
        task_start_patterns = [
            r'^\s*(\d+)\.?\s+',                     # 1. Task description
            r'^\s*Step\s+(\d+)\.?\s+',              # Step 1. Task description
            r'^\s*Task\s+(\d+)\.?\s+',              # Task 1. Task description
            r'^\s*(\d+)\)\s+',                      # 1) Task description
            r'^\s*[Ss][lL]\s*\.?\s*[Nn][oO]\s*\.?\s*(\d+)\s*[\.:]?\s*', # SL NO 1: or SL. NO. 1.
            r'^\s*[Pp]rocedure\s+(\d+)',            # Procedure 1
            r'^\s*[Oo]peration\s+(\d+)',            # Operation 1
            r'^\s*[Ii]tem\s+(\d+)'                  # Item 1
        ]
        
        # Skip possible document title/headers at the beginning
        start_idx = 0
        header_keywords = ['title', 'document', 'assembly', 'procedure', 'work instruction', 'rev', 'revision']
        while start_idx < min(10, len(paragraphs)) and any(keyword in paragraphs[start_idx].text.lower() for keyword in header_keywords):
            start_idx += 1
        
        # Process paragraphs
        for i in range(start_idx, len(paragraphs)):
            para = paragraphs[i]
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
                
            # Check if this paragraph starts a new task
            is_task_start = False
            task_num = None
            
            for pattern in task_start_patterns:
                match = re.match(pattern, text)
                if match:
                    task_num = int(match.group(1))
                    is_task_start = True
                    break
            
            # If we found a new task
            if is_task_start:
                # Save the previous task if there was one
                if current_task is not None:
                    task_index += 1
                    task_number = f"{assembly_id}.0.{str(task_index).zfill(3)}"
                    
                    task = {
                        'task_no': task_number,
                        'type': 'Operation',
                        'eta_sec': '',
                        'description': assembly_name,
                        'activity': current_task_text.strip(),
                        'specification': '',
                        'attachment': ''
                    }
                    tasks.append(task)
                
                # Start new task with the extracted task number
                current_task = task_num
                task_index = task_num  # Use the extracted number directly
                current_task_text = text
                
                # Extract figure references from this paragraph
                extract_figure_references(text, f"{assembly_id}.0.{str(task_num).zfill(3)}", 
                                         image_task_mapping, image_map, figure_start_range, figure_end_range, assembly_id)
            else:
                # Continue with the current task
                if current_task is not None:
                    current_task_text += " " + text
                    
                    # Extract figure references from this paragraph
                    extract_figure_references(text, f"{assembly_id}.0.{str(current_task).zfill(3)}", 
                                             image_task_mapping, image_map, figure_start_range, figure_end_range, assembly_id)
        
        # Save the last task
        if current_task is not None:
            task_index += 1
            task_number = f"{assembly_id}.0.{str(current_task).zfill(3)}"
            
            task = {
                'task_no': task_number,
                'type': 'Operation',
                'eta_sec': '',
                'description': assembly_name,
                'activity': current_task_text.strip(),
                'specification': '',
                'attachment': ''
            }
            tasks.append(task)
        
        # If no tasks were found with the standard patterns, try more aggressive parsing
        if len(tasks) == 0:
            logger.info("No tasks found with standard patterns, trying aggressive paragraph parsing")
            
            # More aggressive parsing - look for any paragraph that might be a task
            potential_tasks = []
            current_text = ""
            
            # Combine paragraphs that appear to be part of the same section
            for i, para in enumerate(paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                
                # Check if this looks like a new task start
                is_new_section = False
                
                # Look for numbering patterns anywhere in the text
                number_match = re.search(r'(\d+)[\.\s\)]+', text)
                
                if number_match:
                    # Check if it's at the start or prominent in the paragraph
                    start_pos = number_match.start()
                    if start_pos < 5 or (text[:start_pos].strip() == ""):
                        is_new_section = True
                        task_num = int(number_match.group(1))
                
                # Also check for action verbs at start which often indicate task steps
                action_verbs = ['install', 'remove', 'connect', 'check', 'ensure', 'verify', 'position', 'align', 'tighten']
                if any(text.lower().startswith(verb) for verb in action_verbs):
                    is_new_section = True
                    # Use sequential numbering for action verb starts
                    task_num = len(potential_tasks) + 1
                
                if is_new_section and current_text:
                    # Save previous section if it exists
                    potential_tasks.append((len(potential_tasks) + 1, current_text))
                    current_text = text
                elif is_new_section:
                    # Start a new section
                    current_text = text
                else:
                    # Continue current section
                    if current_text:
                        current_text += " " + text
                    else:
                        current_text = text
            
            # Don't forget the last section
            if current_text:
                potential_tasks.append((len(potential_tasks) + 1, current_text))
            
            # Convert potential tasks to the standard format
            for idx, (task_num, text) in enumerate(potential_tasks):
                task_number = f"{assembly_id}.0.{str(task_num).zfill(3)}"
                
                task = {
                    'task_no': task_number,
                    'type': 'Operation',
                    'eta_sec': '',
                    'description': assembly_name,
                    'activity': text.strip(),
                    'specification': '',
                    'attachment': ''
                }
                tasks.append(task)
        
        # Extract images from document
        all_images = extract_all_images_from_document(doc)
        
        # Create final image dictionary with proper IDs
        for fig_num, image_id in image_map.items():
            # Find an image to associate with this figure number
            image_idx = min(fig_num - figure_start_range, len(all_images) - 1) if len(all_images) > 0 else -1
            if 0 <= image_idx < len(all_images):
                images[image_id] = {
                    'data': all_images[image_idx]['data'],
                    'extension': all_images[image_idx]['extension'],
                    'figure_number': fig_num
                }
        
        # If we have more images than figure references, add them with sequential IDs
        for i in range(len(image_map), len(all_images)):
            next_num = figure_start_range + i if figure_start_range > 0 else (i + 1)
            image_id = f"{assembly_id}-0-{str(next_num).zfill(3)}"
            images[image_id] = {
                'data': all_images[i]['data'],
                'extension': all_images[i]['extension'],
                'figure_number': next_num
            }
        
        # Now update each task with its attachment information
        for task in tasks:
            task_num = task['task_no']
            if task_num in image_task_mapping:
                task['attachment'] = ', '.join(sorted(list(image_task_mapping[task_num])))
        
        # If we still have no task-image mappings but have both tasks and images,
        # distribute images evenly among tasks
        if len(image_task_mapping) == 0 and len(tasks) > 0 and len(images) > 0:
            logger.info("No explicit figure references found, distributing images among tasks")
            
            # For each task, assign nearby images based on relative position
            task_count = len(tasks)
            image_count = len(images)
            image_keys = list(images.keys())
            
            # Distribute images proportionally
            for i, task in enumerate(tasks):
                # Calculate which images should go to this task
                start_idx = int((i * image_count) / task_count)
                end_idx = int(((i + 1) * image_count) / task_count)
                
                if start_idx < end_idx:  # Ensure task gets at least one image
                    task_images = image_keys[start_idx:end_idx]
                    task['attachment'] = ', '.join(task_images)
        
        # Create DataFrame
        if tasks:
            df = pd.DataFrame(tasks)
            return df, images
        else:
            logger.warning("No tasks extracted from the document")
            return None, {}
            
    except Exception as e:
        logger.error(f"Error in paragraph extraction: {e}")
        return None, {}

def extract_figure_references(text, task_number, image_task_mapping, image_map, figure_start_range, figure_end_range, assembly_id):
    """Extract figure references from text and update mappings"""
    # Look for figure references like "Figure 1", "Fig. 2", "Fig 3", etc.
    figure_patterns = [
        r'[Ff]igure\s+(\d+)',
        r'[Ff]ig\.?\s+(\d+)',
        r'[Ii]llustration\s+(\d+)',
        r'[Pp]hoto\s+(\d+)',
        r'[Pp]icture\s+(\d+)',
        r'[Ii]mage\s+(\d+)',
        r'[Dd]iagram\s+(\d+)',
        r'[Ee]xhibit\s+(\d+)'
    ]
    
    for pattern in figure_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            try:
                fig_num = int(match)
                
                # If no figure range is specified (both are 0), accept all figures
                if figure_start_range == 0 and figure_end_range == 0:
                    accept_figure = True
                else:
                    # Otherwise, check if figure is within range
                    accept_figure = figure_start_range <= fig_num <= figure_end_range
                
                if accept_figure:
                    # Generate image ID (format: assembly_id-0-XXX)
                    image_id = f"{assembly_id}-0-{str(fig_num).zfill(3)}"
                    
                    # Track which figures are referenced by this task
                    if task_number not in image_task_mapping:
                        image_task_mapping[task_number] = set()
                    
                    image_task_mapping[task_number].add(image_id)
                    
                    # Track figure number to image ID mapping
                    image_map[fig_num] = image_id
            except ValueError:
                continue

def extract_all_images_from_document(doc):
    """Extract all images from a Word document with improved methods"""
    all_images = []
    
    # Method 1: Get images from relationships
    try:
        # Get all the relationships in the document
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    all_images.append({
                        'id': rel.rId,
                        'data': image_data,
                        'extension': guess_image_extension(image_data)
                    })
                except Exception as e:
                    logger.error(f"Error extracting image relationship: {e}")
    except Exception as e:
        logger.error(f"Error in relationship extraction: {e}")
    
    # Method 2: Try to directly traverse the document XML structure
    try:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if hasattr(run, '_element') and run._element is not None:
                    for element in run._element.findall('.//{*}drawing') or []:
                        for blip in element.findall('.//{*}blip') or []:
                            embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed and embed in doc.part.rels:
                                try:
                                    rel = doc.part.rels[embed]
                                    image_data = rel.target_part.blob
                                    all_images.append({
                                        'id': embed,
                                        'data': image_data,
                                        'extension': guess_image_extension(image_data)
                                    })
                                except Exception as e:
                                    logger.error(f"Error extracting embedded image: {e}")
    except Exception as e:
        logger.error(f"Error in direct XML traversal: {e}")
    
    # Method 3: Look for inline shapes (another way images can be stored)
    try:
        for paragraph in doc.paragraphs:
            if hasattr(paragraph, 'inline_shapes'):
                for shape in paragraph.inline_shapes:
                    if hasattr(shape, '_inline') and hasattr(shape._inline, 'graphic'):
                        # This is a complex process - log that we found something
                        logger.info(f"Found potential inline shape")
                        # Implementation details would depend on docx structure
    except Exception as e:
        logger.error(f"Error searching for inline shapes: {e}")
    
    # Remove duplicate images by comparing content
    unique_images = []
    seen_hashes = set()
    
    for img in all_images:
        # Use first 500 bytes as a simple hash to identify duplicates
        img_hash = hash(img['data'][:500])
        if img_hash not in seen_hashes:
            seen_hashes.add(img_hash)
            unique_images.append(img)
    
    logger.info(f"Extracted {len(unique_images)} unique images from document")
    return unique_images

def guess_image_extension(image_data):
    """
    Determine the image file extension based on its data.
    """
    import imghdr
    img_type = imghdr.what(None, h=image_data)
    if img_type == 'jpeg':
        return 'jpg'
    return img_type or 'png'

def save_tasks_to_excel(df, assembly_name, output_path):
    """
    Save tasks to Excel file with proper formatting.
    
    Args:
        df: DataFrame containing tasks
        assembly_name: Name of the assembly to set as description
        output_path: Path to save the Excel file
    """
    try:
        # Set the description column to the assembly name
        if df is not None and not df.empty:
            # Make column headers match the expected format
            df = df.rename(columns={
                'task_no': 'task_no',
                'type': 'type',
                'eta_sec': 'eta_sec',
                'description': 'description',
                'activity': 'activity',
                'specification': 'specification',
                'attachment': 'attachment'
            })
            
            # Ensure the description is set to assembly name for all rows
            df['description'] = assembly_name
            
            # Create a new Excel workbook with proper formatting
            wb = Workbook()
            ws = wb.active
            
            # Add headers with formatting
            headers = ['task_no', 'type', 'eta_sec', 'description', 'activity', 'specification', 'attachment']
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                # Apply red font color to first 3 columns
                if col_idx <= 3:
                    cell.font = Font(color="FF0000")
            
            # Add data rows
            for row_idx, row in enumerate(df.itertuples(index=False), 2):
                for col_idx, value in enumerate(row, 1):
                    ws.cell(row=row_idx, column=col_idx, value=value)
            
            # Save to Excel
            wb.save(output_path)
            return True
        return False
    except Exception as e:
        logger.error(f"Error saving Excel file: {e}")
        return False

def save_images(images, output_dir):
    """
    Save extracted images to directory.
    
    Args:
        images: Dictionary mapping image IDs to image data
        output_dir: Directory to save images
    """
    try:
        os.makedirs(output_dir, exist_ok=True)
        for image_id, image_info in images.items():
            output_path = os.path.join(output_dir, f"{image_id}.{image_info['extension']}")
            with open(output_path, 'wb') as f:
                f.write(image_info['data'])
        return True
    except Exception as e:
        logger.error(f"Error saving images: {e}")
        return False

def create_zip_package(excel_path, images_dir, output_path, assembly_name):
    """
    Create a ZIP package containing Excel file and images.
    
    Args:
        excel_path: Path to Excel file
        images_dir: Directory containing images
        output_path: Path to save the ZIP file
        assembly_name: Name of the assembly to use for file naming
    """
    try:
        with zipfile.ZipFile(output_path, 'w') as zipf:
            # Add Excel file with assembly name
            zipf.write(excel_path, arcname=f"{assembly_name}.xlsx")
            
            # Add images
            for root, _, files in os.walk(images_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.join('images', file)
                    zipf.write(file_path, arcname=arcname)
        return True
    except Exception as e:
        logger.error(f"Error creating ZIP package: {e}")
        return False

# Main function to process document
def process_document(input_file, assembly_id, assembly_name, figure_start, figure_end):
    """
    Process a Word document to extract tasks and images.
    
    Args:
        input_file: Path to input Word document
        assembly_id: Assembly sequence ID
        assembly_name: Name of the assembly (description)
        figure_start: Starting figure reference number
        figure_end: Ending figure reference number
    
    Returns:
        Dictionary with processing results
    """
    try:
        # Create temp directory
        temp_dir = os.path.join(os.path.dirname(input_file), "temp_processing")
        os.makedirs(temp_dir, exist_ok=True)
        
        logger.info(f"Processing document with assembly ID: {assembly_id}, name: {assembly_name}")
        logger.info(f"Figure range: {figure_start} to {figure_end}")
        
        # Extract tasks and images
        tasks_df, images = extract_tasks_from_word(
            input_file, 
            assembly_id, 
            assembly_name, 
            figure_start, 
            figure_end
        )
        
        if tasks_df is None or tasks_df.empty:
            return {'success': False, 'message': 'No tasks could be extracted from the document'}
        
        # Save tasks to Excel - use assembly_name for the file name
        excel_path = os.path.join(temp_dir, f"{assembly_name}.xlsx")
        save_tasks_to_excel(tasks_df, assembly_name, excel_path)
        
        # Save images
        images_dir = os.path.join(temp_dir, "images")
        save_images(images, images_dir)
        
        # Create ZIP package - use assembly_name for the zip file name
        zip_path = os.path.join(os.path.dirname(input_file), f"{assembly_name}_extracted_data.zip")
        create_zip_package(excel_path, images_dir, zip_path, assembly_name)
        
        # Return results
        return {
            'success': True,
            'message': f'Successfully processed document. Extracted {len(tasks_df)} tasks and {len(images)} images.',
            'tasks': tasks_df.to_dict('records'),
            'images_count': len(images),
            'excel_path': excel_path,
            'zip_path': zip_path
        }
    except Exception as e:
        logger.error(f"Error in process_document: {str(e)}")
        return {'success': False, 'message': f'Error: {str(e)}'}

# Command line interface
if __name__ == "__main__":
    if len(sys.argv) < 6:
        print("Usage: python script.py input_file assembly_id assembly_name figure_start figure_end")
        sys.exit(1)
    
    input_file = sys.argv[1]
    assembly_id = sys.argv[2]
    assembly_name = sys.argv[3]
    figure_start = int(sys.argv[4])
    figure_end = int(sys.argv[5])
    
    result = process_document(input_file, assembly_id, assembly_name, figure_start, figure_end)
    # Print as JSON for the Node.js wrapper to parse
    print(json.dumps(result))
